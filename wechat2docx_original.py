#!/usr/bin/env python3
"""
wechat2docx.py - 把微信公众号文章链接转换成 Word 文档（保留图片）

用法：
    python3 wechat2docx.py "https://mp.weixin.qq.com/s/xxxx"
    python3 wechat2docx.py "https://mp.weixin.qq.com/s/xxxx" -o 输出.docx
    python3 wechat2docx.py --comments "https://mp.weixin.qq.com/s/xxxx"   # 需要先填好 COOKIES

依赖：
    pip install requests beautifulsoup4 python-docx pillow matplotlib resvg-py
"""

import argparse
import os
import re
import sys
import json
import tempfile
from concurrent.futures import ThreadPoolExecutor, wait as futures_wait, FIRST_COMPLETED
from io import BytesIO
from urllib.parse import urlparse, parse_qs

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from PIL import Image


# ============================================================
# 评论功能配置（仅在使用 --comments 时需要填写）
# ============================================================
# 评论是私有接口，需要从微信 App 抓包后填入以下信息。
# 抓包方法：
#   1. 用 Charles / Fiddler / mitmproxy 在手机上代理
#   2. 在微信里打开目标公众号文章 → 点开"留言/评论"
#   3. 找到形如 https://mp.weixin.qq.com/mp/appmsg_comment?action=getcomment&... 的请求
#   4. 复制其中的 Cookie 头，以及 URL 中的 appmsgid/comment_id/__biz 等参数
#
# 提示：Cookie 有时效，过期需要重新抓取。
COMMENTS_COOKIE = ""    # 形如: "wxuin=xxx; pass_ticket=yyy; ..."
COMMENTS_USER_AGENT = (
    "Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) "
    "AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148 MicroMessenger/8.0.50"
)


# ============================================================
# 抓取与解析
# ============================================================

DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
}


def fetch_article(url: str) -> BeautifulSoup:
    for attempt in range(3):
        try:
            resp = requests.get(url, headers=DEFAULT_HEADERS, timeout=20)
            resp.raise_for_status()
            resp.encoding = resp.apparent_encoding or "utf-8"
            return BeautifulSoup(resp.text, "html.parser")
        except requests.exceptions.SSLError:
            if attempt == 2:
                raise
            import time as _time
            _time.sleep(1)


def extract_meta(soup: BeautifulSoup) -> dict:
    """提取标题、作者、公众号、发布时间"""
    title = ""
    h = soup.select_one("#activity-name, h1.rich_media_title")
    if h:
        title = h.get_text(strip=True)
    if not title:
        og = soup.find("meta", property="og:title")
        if og:
            title = og.get("content", "").strip()

    author = ""
    a = soup.select_one("#js_author_name, .rich_media_meta_text")
    if a:
        author = a.get_text(strip=True)

    account = ""
    acc = soup.select_one("#js_name, .profile_nickname")
    if acc:
        account = acc.get_text(strip=True)

    publish_time = ""
    t = soup.select_one("#publish_time, em#publish_time")
    if t:
        publish_time = t.get_text(strip=True)
    # 如果 DOM 里的时间为空（由 JS 动态填充），从 script 变量 ct 提取
    if not publish_time:
        for script in soup.find_all("script"):
            txt = script.string or ""
            m = re.search(r"var\s+ct\s*=\s*['\"]?(\d{10})['\"]?\s*;", txt)
            if m:
                import datetime
                ts = int(m.group(1))
                publish_time = datetime.datetime.fromtimestamp(ts).strftime("%Y年%-m月%-d日 %H:%M")
                break

    return {
        "title": title or "无标题",
        "author": author,
        "account": account,
        "publish_time": publish_time,
    }


IMG_SIZE_LIMIT = 8 * 1024 * 1024  # 8MB，跳过超大图片

# 全局 Session，复用 TCP 连接，减少握手开销
_img_session = requests.Session()
_img_session.headers.update(DEFAULT_HEADERS)
adapter = requests.adapters.HTTPAdapter(
    pool_connections=10, pool_maxsize=10, max_retries=0
)
_img_session.mount("https://", adapter)
_img_session.mount("http://", adapter)


def download_image(url: str, referer: str = "") -> bytes | None:
    try:
        if url.startswith("//"):
            url = "https:" + url
        headers = {"Referer": referer} if referer else {}
        # timeout=(connect, read)：连接 3s，每次收数据 8s
        resp = _img_session.get(url, headers=headers, timeout=(3, 8), stream=True)
        resp.raise_for_status()
        # 流式读取，防止超大图片卡死
        chunks = []
        size = 0
        for chunk in resp.iter_content(65536):
            size += len(chunk)
            if size > IMG_SIZE_LIMIT:
                print(f"  [图片过大跳过] {url[:80]}", file=sys.stderr)
                return None
            chunks.append(chunk)
        return b"".join(chunks)
    except Exception as e:
        print(f"  [图片下载失败] {url[:80]}... → {e}", file=sys.stderr)
        return None


def _add_table_to_doc(doc: Document, table_elem, set_size_fn) -> None:
    """将 HTML table 转换为 Word 表格"""
    rows = table_elem.find_all("tr")
    if not rows:
        return
    max_cols = max((len(r.find_all(["td", "th"])) for r in rows), default=0)
    if max_cols == 0:
        return
    word_table = doc.add_table(rows=len(rows), cols=max_cols)
    try:
        word_table.style = "Table Grid"
    except Exception:
        pass
    for r_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for c_idx, cell in enumerate(cells[:max_cols]):
            text = cell.get_text(strip=True)
            wc = word_table.cell(r_idx, c_idx)
            wc.text = text
            para = wc.paragraphs[0]
            set_size_fn(para, 10.5)
            if cell.name == "th":
                for run in para.runs:
                    run.bold = True


def add_image_to_doc(doc: Document, img_bytes: bytes):
    """把图片 bytes 加入 Word，自动转换 webp/避免格式问题"""
    try:
        img = Image.open(BytesIO(img_bytes))
        # docx 不支持 webp/gif 动图，统一转 PNG
        if img.format not in ("PNG", "JPEG"):
            buf = BytesIO()
            img.convert("RGB").save(buf, format="PNG")
            buf.seek(0)
        else:
            buf = BytesIO(img_bytes)
        # 限制宽度，避免溢出页面
        doc.add_picture(buf, width=Inches(5.5))
    except Exception as e:
        print(f"  [图片插入失败] {e}", file=sys.stderr)


# ============================================================
# 公式渲染（matplotlib mathtext，无需系统 TeX）
# ============================================================

# 外层公式 span 的特征：style 含 cursor:pointer
_OUTER_FORMULA_RE = re.compile(r"cursor\s*:\s*pointer", re.I)


def mjx_svg_to_png(mjx_elem) -> bytes | None:
    """将 <mjx-container> 内的 SVG 用 resvg-py 渲染为 PNG（透明背景）"""
    try:
        from resvg_py import svg_to_bytes
        svg = mjx_elem.find("svg")
        if not svg:
            return None
        svg_str = str(svg).replace("currentColor", "black")
        if "xmlns=" not in svg_str:
            svg_str = svg_str.replace("<svg", '<svg xmlns="http://www.w3.org/2000/svg"', 1)
        # BeautifulSoup html.parser 把属性全转小写，但 SVG 标准要求驼峰；
        # 不修复的话 resvg 取不到 viewBox/preserveAspectRatio，只渲染默认 0,0 区域 → 碎片
        svg_str = re.sub(r"\bviewbox=", "viewBox=", svg_str)
        svg_str = re.sub(r"\bpreserveaspectratio=", "preserveAspectRatio=", svg_str)
        svg_str = re.sub(r"\bxlink:href=", "xlink:href=", svg_str)  # html.parser 一般不动冒号
        # font_size=24 使 ex 单位对应合理像素大小；zoom=3 提升清晰度
        return svg_to_bytes(svg_str, font_size=24, zoom=3, dpi=0)
    except Exception as e:
        print(f"  [公式SVG渲染失败] {e}", file=sys.stderr)
        return None


def formula_to_png_bytes(latex_code: str) -> bytes | None:
    """将 LaTeX 代码渲染为透明背景 PNG，使用 matplotlib mathtext"""
    try:
        import os as _os
        import tempfile as _tmp
        # 避免 matplotlib 因 sandbox 下配置目录不可写而报错
        _os.environ.setdefault("MPLCONFIGDIR", _tmp.gettempdir())
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig = plt.figure(figsize=(0.01, 0.01))
        fig.patch.set_facecolor("none")
        text = r"$" + latex_code + r"$"
        t = fig.text(0, 0, text, fontsize=11)
        fig.canvas.draw()
        bbox = t.get_window_extent()
        dpi = 150
        w = max(bbox.width / dpi + 0.05, 0.1)
        h = max(bbox.height / dpi + 0.05, 0.1)
        fig.set_size_inches(w, h)
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                    facecolor="none", transparent=True, pad_inches=0.01)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"  [公式渲染失败] {latex_code[:50]}: {e}", file=sys.stderr)
        return None


def _color_str_to_rgb(color_str: str):
    """将 CSS 颜色字符串（#rrggbb / #rgb / rgb(r,g,b)）转为 (r,g,b) 元组，失败返回 None"""
    try:
        s = color_str.strip()
        if s.startswith("#"):
            h = s[1:]
            if len(h) == 3:
                h = h[0]*2 + h[1]*2 + h[2]*2
            if len(h) >= 6:
                return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        m = re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)', s)
        if m:
            return int(m.group(1)), int(m.group(2)), int(m.group(3))
    except Exception:
        pass
    return None


def _parse_css_colors(soup) -> dict:
    """从页面 <style> 标签提取 .class → 前景色字符串 映射（用于代码语法高亮等）"""
    color_map = {}
    for style_tag in soup.find_all("style"):
        css = style_tag.string or ""
        for m in re.finditer(r'\.(\w[\w-]*)[^{]*\{([^}]*)\}', css, re.S):
            cls = m.group(1)
            rules = m.group(2)
            # 只提取 color（不含 background-color）
            cm = re.search(r'(?<![a-z-])color\s*:\s*(#[0-9a-fA-F]{3,8}|rgb\([^)]+\))', rules)
            if cm:
                color_map[cls] = cm.group(1)
    return color_map


# CSS color inline style 提取正则
_CSS_COLOR_RE = re.compile(r'(?<![a-z-])color\s*:\s*(#[0-9a-fA-F]{3,8}|rgb\([^)]+\))', re.I)

# 微信 / mdnice 代码高亮类名 → 颜色（静态备用，CSS 外链无法抓取时使用）
# 参考 VS Code Light+ 主题，与微信编辑器默认高亮接近
_WECHAT_CODE_COLORS = {
    "code-snippet__keyword":    "#0000ff",
    "code-snippet__operator":   "#0000ff",
    "code-snippet__builtin":    "#0070c1",
    "code-snippet__built_in":   "#0070c1",
    "code-snippet__string":     "#a31515",
    "code-snippet__comment":    "#008000",
    "code-snippet__number":     "#098658",
    "code-snippet__literal":    "#098658",
    "code-snippet__class-name": "#267f99",
    "code-snippet__function":   "#795e26",
    "code-snippet__attribute":  "#001080",
    "code-snippet__property":   "#001080",
    "code-snippet__variable":   "#001080",
    "code-snippet__params":     "#001080",
    "code-snippet__tag":        "#800000",
    "code-snippet__type":       "#267f99",
    "code-snippet__symbol":     "#a31515",
    "code-snippet__selector":   "#800000",
    "code-snippet__meta":       "#808080",
    "code-snippet__deletion":   "#9a1212",
    "code-snippet__addition":   "#008000",
    # hljs 类名（部分文章使用 hljs 高亮器）
    "hljs-keyword":   "#0000ff",
    "hljs-string":    "#a31515",
    "hljs-comment":   "#008000",
    "hljs-number":    "#098658",
    "hljs-built_in":  "#0070c1",
    "hljs-attr":      "#001080",
    "hljs-function":  "#795e26",
    "hljs-title":     "#795e26",
    "hljs-type":      "#267f99",
    "hljs-meta":      "#808080",
    "hljs-literal":   "#098658",
}


def _populate_para_with_colors(p, elem, css_colors=None):
    """向已有段落添加带颜色的 run（支持行内公式、语法高亮色、内联 style 颜色）"""
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor

    def _walk(node, color=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).replace("\xa0", " ")
                if text.strip():
                    run = p.add_run(text)
                    if color:
                        run.font.color.rgb = RGBColor(*color)
            elif isinstance(child, Tag):
                # 外层公式 span（cursor:pointer + data-formula）
                if (child.name == "span"
                        and child.has_attr("data-formula")
                        and _OUTER_FORMULA_RE.search(child.get("style", ""))):
                    latex = child["data-formula"]
                    png = formula_to_png_bytes(latex)
                    if png:
                        p.add_run().add_picture(BytesIO(png), height=Pt(10.5))
                    else:
                        p.add_run(f"[{latex}]")
                # 行内 MathJax 容器（无 display="true"，仅含 SVG）
                elif child.name == "mjx-container" and not child.get("display"):
                    png = mjx_svg_to_png(child)
                    if png:
                        p.add_run().add_picture(BytesIO(png), height=Pt(10.5))
                elif child.name == "br":
                    p.add_run("\n")
                elif child.name == "img":
                    pass  # 图片由主循环处理
                elif child.name in ("svg", "mjx-container"):
                    pass  # SVG 和块级 mjx 由上层处理，不递归
                else:
                    # 提取颜色：优先内联 style，其次 CSS 类
                    new_color = color
                    cm = _CSS_COLOR_RE.search(child.get("style", ""))
                    if cm:
                        new_color = _color_str_to_rgb(cm.group(1)) or color
                    if new_color is color and css_colors:
                        for cls in (child.get("class") or []):
                            if cls in css_colors:
                                new_color = _color_str_to_rgb(css_colors[cls]) or color
                                break
                    _walk(child, new_color)

    _walk(elem)


def _add_rich_para(doc, elem, pt, align=None, css_colors=None):
    """创建段落，将文本 run 和行内公式图片 run 混合排列，保留原始颜色"""
    p = doc.add_paragraph()
    _populate_para_with_colors(p, elem, css_colors)
    _set_para_size(p, pt, align)
    return p


def _set_para_size(para, pt, align=None):
    """为段落设置字号、段前段后0、单倍行距，可选对齐方式"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # 段落级别默认字号（pPr/rPr）
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(pt * 2)))
    # 各 run 字号
    for run in para.runs:
        run.font.size = Pt(pt)
    # 段间距：段前段后0，单倍行距
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align is not None:
        para.alignment = align


def _add_code_para_colored(doc, pre_elem, css_colors=None):
    """将 <pre> 代码块转为带语法高亮颜色的 Word 段落"""
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor

    p = doc.add_paragraph()
    for s in doc.styles:
        if s.name == "No Spacing":
            p.style = s
            break

    def _add_run(text, color=None):
        if not text:
            return
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _walk_code(node, color=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                _add_run(str(child).replace("\xa0", " "), color)
            elif isinstance(child, Tag):
                new_color = color
                # 内联 style
                cm = _CSS_COLOR_RE.search(child.get("style", ""))
                if cm:
                    new_color = _color_str_to_rgb(cm.group(1)) or color
                # CSS 类（代码高亮）：先查动态 CSS，再 fallback 到 _WECHAT_CODE_COLORS
                if new_color is color:
                    for cls in (child.get("class") or []):
                        if css_colors and cls in css_colors:
                            new_color = _color_str_to_rgb(css_colors[cls]) or color
                            break
                        if cls in _WECHAT_CODE_COLORS:
                            new_color = _color_str_to_rgb(_WECHAT_CODE_COLORS[cls]) or color
                            break
                _walk_code(child, new_color)

    code_children = pre_elem.find_all("code")
    if code_children:
        for i, code in enumerate(code_children):
            if i > 0:
                _add_run("\n")
            _walk_code(code)
    else:
        # 降级：无 <code> 子元素，直接按 <br> 分行
        for br in pre_elem.find_all("br"):
            br.replace_with("\n")
        _add_run(pre_elem.get_text().replace("\xa0", " ").strip("\n").rstrip())

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def build_docx(soup: BeautifulSoup, meta: dict, output: str, comments: list = None):
    doc = Document()

    # 标题：去掉末尾标点，字号四号（14pt），居中，段前段后0
    clean_title = meta["title"].rstrip().rstrip("。！？，、；：…—,.!?;:")
    heading = doc.add_heading(clean_title, level=0)
    _set_para_size(heading, 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 元信息（五号 10.5pt）
    info_parts = [p for p in [meta["account"], meta["author"], meta["publish_time"]] if p]
    if info_parts:
        p = doc.add_paragraph(" | ".join(info_parts))
        _set_para_size(p, 10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 空行（段间距也清零）
    blank = doc.add_paragraph()
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(0)
    blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 正文
    content = soup.select_one("#js_content, .rich_media_content")
    if not content:
        body_text = soup.body.get_text(strip=True) if soup.body else ""
        if "轻触阅读原文" in body_text or "阅读原文" in body_text:
            reason = "外链跳转卡片（内容在外部网站，无正文可抓取）"
        else:
            reason = "付费/受限/仅粉丝可见文章"
        print(f"[警告] 未找到正文容器：{reason}", file=sys.stderr)
        raise ValueError(f"无法提取正文：{reason}")

    # 提取页面 CSS 颜色映射（用于代码语法高亮、彩色文本等）
    css_colors = _parse_css_colors(soup)

    # 并发预下载所有图片
    img_cache: dict = {}
    img_elems = [(e, e.get("data-src") or e.get("src")) for e in content.find_all("img")]
    img_elems = [(e, u) for e, u in img_elems if u]
    # 从 meta 或 og:url 取文章 URL 作 Referer（帮助绕过 CDN 鉴权）
    referer = ""
    og = soup.find("meta", property="og:url")
    if og:
        referer = og.get("content", "")
    if img_elems:
        print(f"  并发下载 {len(img_elems)} 张图片…", file=sys.stderr)
        with ThreadPoolExecutor(max_workers=10) as pool:
            future_map = {
                pool.submit(download_image, url, referer): elem
                for elem, url in img_elems
            }
            # 全局最多等 30s，超时直接放弃剩余图片
            done, not_done = futures_wait(future_map, timeout=30)
            for future in not_done:
                future.cancel()
                print("  [图片超时跳过]", file=sys.stderr)
            for future in done:
                elem = future_map[future]
                try:
                    img_cache[id(elem)] = future.result()
                except Exception:
                    pass

    img_count = 0
    para_count = 0
    skip_elems: set = set()  # 已由父元素处理、需跳过的元素 id

    for elem in content.descendants:
        if id(elem) in skip_elems:
            continue
        # 跳过 table 内部元素，避免重复（由 _add_table_to_doc 统一处理）
        if elem.find_parent("table"):
            continue
        # 跳过 pre 内部元素，避免重复（由 pre 分支统一处理）
        if elem.find_parent("pre"):
            continue
        # 跳过公式 span 内部元素（SVG 等，由 _add_rich_para 统一处理）
        if elem.find_parent("span", attrs={"data-formula": True}):
            continue
        # 跳过 mjx-container 内部元素（SVG，由各自分支统一处理）
        if elem.find_parent("mjx-container"):
            continue
        # 块级 MathJax 公式（display=true），渲染为图片独占一行
        if elem.name == "mjx-container" and elem.get("display") == "true":
            png = mjx_svg_to_png(elem)
            if png:
                try:
                    doc.add_picture(BytesIO(png), width=Inches(5.5))
                    para_count += 1
                except Exception as e:
                    print(f"  [块级公式插入失败] {e}", file=sys.stderr)
            continue
        if elem.name == "table":
            _add_table_to_doc(doc, elem, _set_para_size)
            para_count += 1
        elif elem.name == "pre":
            # 代码块：使用 _add_code_para_colored 保留语法高亮颜色
            p = _add_code_para_colored(doc, elem, css_colors)
            # 检查是否有内容（纯空白 pre 跳过）
            if not p.text.strip():
                p._element.getparent().remove(p._element)
            else:
                para_count += 1
        elif elem.name == "img":
            img_bytes = img_cache.get(id(elem))
            if img_bytes:
                add_image_to_doc(doc, img_bytes)
                img_count += 1
        elif elem.name in ("p", "section", "h1", "h2", "h3", "h4", "li", "blockquote"):
            # 只在叶子段落写文字，避免重复
            if elem.find(["p", "section", "li"]):
                continue

            # 检测行内公式：外层 span(data-formula+cursor:pointer) 或行内 mjx-container
            outer_formulas = elem.find_all(
                "span", attrs={"data-formula": True, "style": _OUTER_FORMULA_RE}
            )
            inline_mjx = elem.find_all("mjx-container", attrs={"display": False})
            # find_all 返回所有 mjx-container，再过滤掉 display="true" 的块级
            inline_mjx = [m for m in elem.find_all("mjx-container") if not m.get("display")]
            if outer_formulas or inline_mjx:
                # 将外层公式 span 及其子元素标记为已处理
                for fs in outer_formulas:
                    skip_elems.add(id(fs))
                    for d in fs.descendants:
                        skip_elems.add(id(d))
                # 将行内 mjx-container 及其子元素标记为已处理（_add_rich_para 会处理它们）
                for mj in inline_mjx:
                    skip_elems.add(id(mj))
                    for d in mj.descendants:
                        skip_elems.add(id(d))
                if elem.name in ("h1", "h2"):
                    p = _add_rich_para(doc, elem, 12, css_colors=css_colors)
                elif elem.name in ("h3", "h4"):
                    p = _add_rich_para(doc, elem, 10.5, css_colors=css_colors)
                elif elem.name == "blockquote":
                    p = _add_rich_para(doc, elem, 10.5, css_colors=css_colors)
                    try:
                        p.style = "Intense Quote"
                    except Exception:
                        pass
                else:
                    p = _add_rich_para(doc, elem, 10.5, css_colors=css_colors)
                para_count += 1
                continue

            # 普通段落：也用 _add_rich_para 以保留内联颜色
            if elem.name in ("h1", "h2"):
                p = doc.add_heading("", level=2)
                _populate_para_with_colors(p, elem, css_colors)
                _set_para_size(p, 12)  # 小四
            elif elem.name in ("h3", "h4"):
                p = doc.add_heading("", level=3)
                _populate_para_with_colors(p, elem, css_colors)
                _set_para_size(p, 10.5)  # 五号
            elif elem.name == "blockquote":
                p = _add_rich_para(doc, elem, 10.5, css_colors=css_colors)
                p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else p.style
            else:
                p = _add_rich_para(doc, elem, 10.5, css_colors=css_colors)
            if not p.text.strip():
                continue
            para_count += 1

    print(f"  正文：{para_count} 段落，{img_count} 张图片", file=sys.stderr)

    # 评论
    if comments:
        doc.add_page_break()
        doc.add_heading("评论", level=1)
        for c in comments:
            nick = c.get("nick_name", "匿名")
            content_txt = c.get("content", "")
            like = c.get("like_num", 0)
            doc.add_paragraph(f"【{nick}】（赞 {like}）", style="Intense Quote" if False else None)
            doc.add_paragraph(content_txt)
            for reply in c.get("reply", {}).get("reply_list", []):
                doc.add_paragraph(
                    f"    ↳ {reply.get('nick_name','')}：{reply.get('content','')}"
                )
        print(f"  评论：{len(comments)} 条", file=sys.stderr)

    doc.save(output)
    print(f"  → {output}", file=sys.stderr)


# ============================================================
# 评论抓取（私有接口）
# ============================================================

def extract_comment_params(html: str) -> dict:
    """从文章 HTML 里抠出评论需要的参数"""
    params = {}
    for key in ("biz", "appmsgid", "comment_id", "idx", "sn"):
        m = re.search(rf'var\s+{key}\s*=\s*["\']([^"\']+)["\']', html)
        if m:
            params[key] = m.group(1)
    # __biz 在 URL / meta 里
    m = re.search(r'var\s+__biz\s*=\s*["\']([^"\']+)["\']', html)
    if m:
        params["__biz"] = m.group(1)
    return params


def fetch_comments(article_url: str) -> list:
    if not COMMENTS_COOKIE:
        print(
            "[评论] 未配置 COOKIES，跳过。请编辑脚本顶部 COMMENTS_COOKIE 变量。",
            file=sys.stderr,
        )
        return []

    headers = {
        "User-Agent": COMMENTS_USER_AGENT,
        "Cookie": COMMENTS_COOKIE,
        "Referer": article_url,
    }
    html = requests.get(article_url, headers=headers, timeout=20).text
    p = extract_comment_params(html)
    if not all(k in p for k in ("__biz", "appmsgid", "comment_id", "idx")):
        print(f"[评论] 无法解析参数：{p}", file=sys.stderr)
        return []

    api = "https://mp.weixin.qq.com/mp/appmsg_comment"
    qs = {
        "action": "getcomment",
        "scene": "0",
        "__biz": p["__biz"],
        "appmsgid": p["appmsgid"],
        "idx": p["idx"],
        "comment_id": p["comment_id"],
        "offset": "0",
        "limit": "100",
        "f": "json",
    }
    try:
        r = requests.get(api, params=qs, headers=headers, timeout=20)
        data = r.json()
        return data.get("elected_comment", []) + data.get("friend_comment", [])
    except Exception as e:
        print(f"[评论] 抓取失败：{e}", file=sys.stderr)
        return []


# ============================================================
# CLI
# ============================================================

def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name).strip()[:80] or "wechat_article"


def main():
    parser = argparse.ArgumentParser(description="微信公众号文章 → Word 文档")
    parser.add_argument("url", help="微信文章链接 https://mp.weixin.qq.com/s/...")
    parser.add_argument("-o", "--output", help="输出 docx 路径（默认用文章标题）")
    parser.add_argument("--comments", action="store_true", help="抓取评论（需先配置 COOKIES）")
    args = parser.parse_args()

    print(f"抓取：{args.url}", file=sys.stderr)
    soup = fetch_article(args.url)
    meta = extract_meta(soup)
    print(f"标题：{meta['title']}", file=sys.stderr)

    output = args.output or f"{safe_filename(meta['title'])}.docx"

    comments = fetch_comments(args.url) if args.comments else None
    build_docx(soup, meta, output, comments=comments)


if __name__ == "__main__":
    main()
