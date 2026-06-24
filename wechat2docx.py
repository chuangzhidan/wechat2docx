#!/usr/bin/env python3
"""
wechat2docx.py - 把微信公众号文章链接转换成 Word 文档（保留图片）

用法：
    python3 wechat2docx.py "https://mp.weixin.qq.com/s/xxxx"
    python3 wechat2docx.py "https://mp.weixin.qq.com/s/xxxx" -o 输出.docx
    python3 wechat2docx.py --comments "https://mp.weixin.qq.com/s/xxxx"   # 需要先填好 COOKIES

依赖：
    pip install requests beautifulsoup4 python-docx pillow matplotlib resvg-py
"""

import argparse
import html as html_lib
import os
import re
import sys
import json
import tempfile
from contextvars import ContextVar
from concurrent.futures import ThreadPoolExecutor, wait as futures_wait, FIRST_COMPLETED
from io import BytesIO
from urllib.parse import urlparse, parse_qs, parse_qsl, urlencode, urlunparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from PIL import Image


# ============================================================
# 评论功能配置（仅在使用 --comments 时需要填写）
# ============================================================
# 评论是私有接口，需要从微信 App 抓包后填入以下信息。
# 抓包方法：
#   1. 用 Charles / Fiddler / mitmproxy 在手机上代理
#   2. 在微信里打开目标公众号文章 → 点开"留言/评论"
#   3. 找到形如 https://mp.weixin.qq.com/mp/appmsg_comment?action=getcomment&... 的请求
#   4. 复制其中的 Cookie 头，以及 URL 中的 appmsgid/comment_id/__biz 等参数
#
# 提示：Cookie 有时效，过期需要重新抓取。
COMMENTS_COOKIE = ""    # 形如: "wxuin=xxx; pass_ticket=yyy; ..."
COMMENTS_USER_AGENT = (
    "Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) "
    "AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148 MicroMessenger/8.0.50"
)


# ============================================================
# 抓取与解析
# ============================================================

DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
}

SUPPORTED_BODY_FONTS = ("宋体", "宋体-简", "Times New Roman", "楷体", "方正仿宋_GBK")
DEFAULT_BODY_FONT = "宋体-简"
DEFAULT_BODY_FONT_SIZE_PT = 10.5

BODY_FONT = DEFAULT_BODY_FONT
EMOJI_FONT = "Segoe UI Emoji"
KEYCAP_EMOJI_FONT = "Apple Color Emoji"

_BODY_FONT_VAR = ContextVar("wechat2docx_body_font", default=DEFAULT_BODY_FONT)
_BODY_FONT_SIZE_VAR = ContextVar("wechat2docx_body_font_size", default=DEFAULT_BODY_FONT_SIZE_PT)


def _body_font() -> str:
    return _BODY_FONT_VAR.get()


def _body_font_size() -> float:
    return _BODY_FONT_SIZE_VAR.get()


def _normalize_body_font(font_name: str | None) -> str:
    if not font_name:
        return DEFAULT_BODY_FONT
    font_name = font_name.strip()
    if font_name not in SUPPORTED_BODY_FONTS:
        raise ValueError(f"不支持的导出字体：{font_name}")
    return font_name


def _normalize_body_font_size(font_size) -> float:
    if font_size is None or font_size == "":
        return DEFAULT_BODY_FONT_SIZE_PT
    try:
        value = float(font_size)
    except (TypeError, ValueError):
        raise ValueError(f"字号必须是数字：{font_size}")
    if not 6 <= value <= 24:
        raise ValueError("字号必须在 6 到 24 pt 之间")
    return value


def _set_rfonts(rPr, font_name: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), font_name)


def _set_run_font(run, font_name: str):
    run.font.name = font_name
    _set_rfonts(run._element.get_or_add_rPr(), font_name)


# 微信分享链接里常见的跟踪参数，会触发「访问验证」反爬页面，抓取前需剔除
_WECHAT_TRACKING_PARAMS = {
    "click_id", "poc_token", "scene", "from", "isappinstalled",
    "clicktime", "enterid", "ascene", "devicetype", "version",
    "nettype", "abtest_cookie", "key", "uin", "pass_ticket", "wx_header",
}


def _clean_wechat_url(url: str) -> str:
    """剔除微信文章 URL 中触发反爬验证页的跟踪参数。

    /s/<token> 形式：query 全是跟踪参数，整体丢弃。
    /s?__biz=...&mid=...&sn=... 形式：query 是定位文章的必要参数，仅剔除已知跟踪参数。
    """
    parsed = urlparse(url)
    if "mp.weixin.qq.com" not in parsed.netloc:
        return url
    # 路径形式 /s/<token>：query 仅为跟踪参数，直接清空
    if parsed.path.startswith("/s/") and len(parsed.path) > len("/s/"):
        return urlunparse(parsed._replace(query="", fragment=""))
    # 查询形式 /s?__biz=...：保留必要参数，仅去掉已知跟踪参数
    kept = [(k, v) for k, v in parse_qsl(parsed.query, keep_blank_values=True)
            if k not in _WECHAT_TRACKING_PARAMS]
    return urlunparse(parsed._replace(query=urlencode(kept), fragment=""))


def fetch_article(url: str) -> BeautifulSoup:
    """抓取文章页面，自动识别微信/CSDN，并把 URL 存到 soup._url"""
    # 雪球：所有接口被阿里云 WAF JS Challenge 保护，requests 无法通过
    if 'xueqiu.com' in url:
        raise ValueError("雪球暂不支持：其页面受 WAF JS 挑战保护，需要浏览器环境才能访问，无法用 requests 抓取")
    url = _clean_wechat_url(url)
    is_csdn = 'blog.csdn.net' in url
    headers = dict(DEFAULT_HEADERS)
    if is_csdn:
        headers['User-Agent'] = 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    for attempt in range(3):
        try:
            resp = requests.get(url, headers=headers, timeout=20)
            resp.raise_for_status()
            resp.encoding = resp.apparent_encoding or "utf-8"
            soup = BeautifulSoup(resp.text, "html.parser")
            soup._url = url  # 保存 URL 供后续提取作者等信息
            return soup
        except requests.exceptions.SSLError:
            if attempt == 2:
                raise
            import time as _time
            _time.sleep(1)


def extract_meta(soup: BeautifulSoup) -> dict:
    """提取标题、作者、公众号、发布时间（自动识别微信/CSDN）"""
    # 检测是否为 CSDN 页面
    if soup.select_one('.article_content, #article_content, #content_views'):
        return _extract_meta_csdn(soup)
    return _extract_meta_wechat(soup)


def _extract_meta_wechat(soup: BeautifulSoup) -> dict:
    """微信文章元信息"""
    title = ""
    h = soup.select_one("#activity-name, h1.rich_media_title")
    if h:
        title = h.get_text(strip=True)
    if not title:
        og = soup.find("meta", property="og:title")
        if og:
            title = og.get("content", "").strip()

    author = ""
    a = soup.select_one("#js_author_name, .rich_media_meta_text")
    if a:
        author = a.get_text(strip=True)

    account = ""
    acc = soup.select_one("#js_name, .profile_nickname")
    if acc:
        account = acc.get_text(strip=True)
    if not account:
        author_meta = soup.find("meta", attrs={"name": "author"})
        if author_meta:
            account = author_meta.get("content", "").strip()

    publish_time = ""
    t = soup.select_one("#publish_time, em#publish_time")
    if t:
        publish_time = t.get_text(strip=True)
    # 如果 DOM 里的时间为空（由 JS 动态填充），从 script 变量 ct 提取
    if not publish_time:
        for script in soup.find_all("script"):
            txt = script.string or ""
            m = re.search(r"var\s+ct\s*=\s*['\"]?(\d{10})['\"]?\s*;", txt)
            if m:
                import datetime
                ts = int(m.group(1))
                publish_time = datetime.datetime.fromtimestamp(ts).strftime("%Y年%-m月%-d日 %H:%M")
                break

    return {
        "title": title or "无标题",
        "author": author,
        "account": account,
        "publish_time": publish_time,
    }


def _decode_wechat_meta_text(text: str) -> str:
    """解码微信 meta/script 里的正文摘要，保留换行。"""
    if not text:
        return ""

    def _replace_hex(match):
        return chr(int(match.group(1), 16))

    text = re.sub(r"\\x([0-9a-fA-F]{2})", _replace_hex, text)
    text = re.sub(r"\\u([0-9a-fA-F]{4})", _replace_hex, text)
    text = text.replace("\\n", "\n").replace("\\r", "\n")
    for _ in range(2):
        text = html_lib.unescape(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(line.rstrip() for line in text.split("\n")).strip()


def _extract_meta_description_text(soup: BeautifulSoup) -> str:
    """微信图片合集等无正文容器页面，正文常藏在 description / og:description。"""
    candidates = []
    for attrs in (
            {"name": "description"},
            {"property": "og:description"},
            {"name": "twitter:description"}):
        meta = soup.find("meta", attrs=attrs)
        if meta:
            candidates.append(meta.get("content", ""))
    best = ""
    for item in candidates:
        decoded = _decode_wechat_meta_text(item)
        if len(decoded) > len(best):
            best = decoded
    return best


def _extract_meta_csdn(soup: BeautifulSoup) -> dict:
    """CSDN 文章元信息"""
    import re
    # 标题
    title = ""
    for sel in ['h1', '.article-title', '#articleContentId', 'title']:
        elem = soup.select_one(sel)
        if elem:
            title = elem.get_text(strip=True)
            break

    # 作者：直接从 soup._url 提取（最可靠）
    author = ""
    article_url = getattr(soup, '_url', '')
    if article_url:
        m = re.search(r'blog\.csdn\.net/([^/]+)', article_url)
        if m:
            author = m.group(1)
    # fallback：从 og:url 或页面元素提取
    if not author:
        url_meta = soup.find('meta', property='og:url')
        if url_meta:
            url = url_meta.get('content', '')
            m = re.search(r'blog\.csdn\.net/([^/]+)', url)
            if m:
                author = m.group(1)
    if not author:
        for sel in ['.profile-user-name a', '.author-name', '.follow-nickName']:
            elem = soup.select_one(sel)
            if elem:
                author = elem.get_text(strip=True)
                break

    # 发布时间
    publish_time = ""
    # 从 script 标签找
    for script in soup.find_all('script'):
        if not script.string:
            continue
        for m in re.finditer(r'"(?:publishTime|createTime|create_time)"\s*:\s*"([^"]+)"', script.string):
            publish_time = m.group(1)
            break
        if publish_time:
            break
    # 从页面文本找时间
    if not publish_time:
        for p in soup.find_all(['span', 'div', 'time']):
            text = p.get_text()
            m = re.search(r'(\d{4}-\d{2}-\d{2}(?:\s+\d{2}:\d{2}:\d{2})?)', text)
            if m:
                publish_time = m.group(1)
                break

    return {
        "title": title or "无标题",
        "author": author,
        "account": "CSDN",
        "publish_time": publish_time,
    }


IMG_SIZE_LIMIT = 8 * 1024 * 1024  # 8MB，跳过超大图片

# 全局 Session，复用 TCP 连接，减少握手开销
_img_session = requests.Session()
_img_session.headers.update(DEFAULT_HEADERS)
adapter = requests.adapters.HTTPAdapter(
    pool_connections=10, pool_maxsize=10, max_retries=0
)
_img_session.mount("https://", adapter)
_img_session.mount("http://", adapter)


def download_image(url: str, referer: str = "") -> bytes | None:
    try:
        if url.startswith("//"):
            url = "https:" + url
        headers = {"Referer": referer} if referer else {}
        # timeout=(connect, read)：连接 3s，每次收数据 8s
        resp = _img_session.get(url, headers=headers, timeout=(3, 8), stream=True)
        resp.raise_for_status()
        # 流式读取，防止超大图片卡死
        chunks = []
        size = 0
        for chunk in resp.iter_content(65536):
            size += len(chunk)
            if size > IMG_SIZE_LIMIT:
                print(f"  [图片过大跳过] {url[:80]}", file=sys.stderr)
                return None
            chunks.append(chunk)
        return b"".join(chunks)
    except Exception as e:
        print(f"  [图片下载失败] {url[:80]}... → {e}", file=sys.stderr)
        return None


def _is_code_table(table_elem) -> bool:
    """检测 table 是否是微信/mdnice 代码块表格（td > section > code 结构）"""
    rows = table_elem.find_all("tr")
    if not rows:
        return False
    # 取第一个非空 td
    first_td = None
    for tr in rows:
        tds = tr.find_all(["td", "th"])
        if tds:
            first_td = tds[0]
            break
    if first_td is None:
        return False
    # 判断：td 内直接子 section 含有 code 元素
    sections = first_td.find_all("section", recursive=False)
    if not sections:
        return False
    return any(s.find("code") for s in sections)


def _add_code_table_to_doc(doc, table_elem, css_colors=None) -> None:
    """将微信代码块表格（td > section*行 > code*片段）渲染为带语法高亮的 Word 段落"""
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor

    # 找到唯一的代码 td（取第一个含 section>code 的 td）
    code_td = None
    for tr in table_elem.find_all("tr"):
        for td in tr.find_all(["td", "th"]):
            if td.find_all("section", recursive=False):
                code_td = td
                break
        if code_td:
            break
    if not code_td:
        return

    p = doc.add_paragraph()
    for s in doc.styles:
        if s.name == "No Spacing":
            p.style = s
            break

    def _add_run(text, color=None):
        if not text:
            return
        run = p.add_run(text)
        _set_run_font(run, _body_font())
        run.font.size = Pt(_body_font_size())
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _get_color(node):
        """从节点的 inline style 或 CSS 类提取颜色，近白色丢弃"""
        color = None
        cm = _CSS_COLOR_RE.search(node.get("style", ""))
        if cm:
            color = _color_str_to_rgb(cm.group(1))
        if color is None and css_colors:
            for cls in (node.get("class") or []):
                if cls in css_colors:
                    color = _color_str_to_rgb(css_colors[cls])
                    break
                if cls in _WECHAT_CODE_COLORS:
                    color = _color_str_to_rgb(_WECHAT_CODE_COLORS[cls])
                    break
        if color and color[0] > 240 and color[1] > 240 and color[2] > 240:
            color = None
        return color

    def _walk_node(node, color=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).replace("\xa0", " ")
                if text:
                    _add_run(text, color)
            elif isinstance(child, Tag):
                if child.name == "br":
                    _add_run("\n")
                    continue
                new_color = _get_color(child) or color
                _walk_node(child, new_color)

    sections = code_td.find_all("section", recursive=False)
    for i, section in enumerate(sections):
        if i > 0:
            _add_run("\n")
        _walk_node(section)


def _add_table_to_doc(doc: Document, table_elem, set_size_fn) -> None:
    """将 HTML table 转换为 Word 表格"""
    rows = table_elem.find_all("tr")
    if not rows:
        return
    max_cols = max((len(r.find_all(["td", "th"])) for r in rows), default=0)
    if max_cols == 0:
        return
    word_table = doc.add_table(rows=len(rows), cols=max_cols)
    try:
        word_table.style = "Table Grid"
    except Exception:
        pass
    for r_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for c_idx, cell in enumerate(cells[:max_cols]):
            text = cell.get_text(strip=True)
            wc = word_table.cell(r_idx, c_idx)
            wc.text = text
            para = wc.paragraphs[0]
            set_size_fn(para, _body_font_size())
            if cell.name == "th":
                for run in para.runs:
                    run.bold = True


def add_image_to_doc(doc: Document, img_bytes: bytes):
    """把图片 bytes 加入 Word，自动转换 webp/避免格式问题"""
    try:
        img = Image.open(BytesIO(img_bytes))
        # docx 不支持 webp/gif 动图，统一转 PNG
        if img.format not in ("PNG", "JPEG"):
            buf = BytesIO()
            img.convert("RGB").save(buf, format="PNG")
            buf.seek(0)
        else:
            buf = BytesIO(img_bytes)
        # 限制宽度，避免溢出页面
        doc.add_picture(buf, width=Inches(5.5))
    except Exception as e:
        print(f"  [图片插入失败] {e}", file=sys.stderr)


# ============================================================
# 公式渲染（matplotlib mathtext，无需系统 TeX）
# ============================================================

# 外层公式 span 的特征：style 含 cursor:pointer
_OUTER_FORMULA_RE = re.compile(r"cursor\s*:\s*pointer", re.I)


def _is_emoji_cp(cp: int) -> bool:
    """判断 Unicode 码点是否为 emoji / variation selector / ZWJ"""
    return (
        0x2300 <= cp <= 0x27BF or    # Misc Technical (⏳ U+23F3) + Symbols (⚠) + Dingbats (✅✨)
        0x2B00 <= cp <= 0x2BFF or    # Misc Symbols and Arrows (⭐ U+2B50)
        0x1F000 <= cp <= 0x1FAFF or  # Emoji pictographs, transport, misc
        cp == 0x20E3 or              # Combining Enclosing Keycap (⃣)
        0xFE00 <= cp <= 0xFE0F or    # Variation Selectors (⚙️ 中的 \uFE0F)
        cp == 0x200D                 # Zero Width Joiner
    )


# 这些码点是 emoji 序列的组合符/修饰符，前面的非 emoji 字符（如数字）是 base
_EMOJI_MODIFIERS = frozenset([0xFE0F, 0x20E3, 0x200D])


def _split_text_by_emoji(text: str):
    """将文本按 emoji/非 emoji 分段，返回 [(is_emoji: bool, segment: str)] 列表
    
    特殊处理 keycap 序列（1️⃣ = '1' U+FE0F U+20E3）：
    若非 emoji 字符后紧跟 emoji 修饰符（U+FE0F/U+20E3/U+200D），将其视为 emoji base。
    """
    if not text:
        return []
    segments = []
    i = 0
    cur_seg = ""
    cur_is_emoji = None
    while i < len(text):
        c = text[i]
        cp = ord(c)
        is_emoji = _is_emoji_cp(cp)
        # 前瞻：非 emoji 字符后面紧跟 emoji 修饰符时，视为 emoji 序列的 base
        if not is_emoji and i + 1 < len(text) and ord(text[i + 1]) in _EMOJI_MODIFIERS:
            is_emoji = True
        if cur_is_emoji is None:
            cur_is_emoji = is_emoji
        if is_emoji == cur_is_emoji:
            cur_seg += c
        else:
            segments.append((cur_is_emoji, cur_seg))
            cur_seg = c
            cur_is_emoji = is_emoji
        i += 1
    if cur_seg:
        segments.append((cur_is_emoji, cur_seg))
    return segments


def _is_keycap_seg(seg: str) -> bool:
    """判断一段文本是否包含 keycap 序列（[0-9#*] + 可选FE0F + U+20E3）
    这类序列需要由系统 emoji font（Apple Color Emoji / Segoe UI Emoji）整体渲染，
    不能强制设置 w:ascii 字体，否则 Mac 上 digit 和 U+20E3 分开显示为 □。
    """
    return '\u20e3' in seg


def mjx_svg_to_png(mjx_elem) -> bytes | None:
    """将 <mjx-container> 内的 SVG 用 resvg-py 渲染为 PNG（透明背景）"""
    try:
        from resvg_py import svg_to_bytes
        svg = mjx_elem.find("svg")
        if not svg:
            return None
        svg_str = str(svg).replace("currentColor", "black")
        if "xmlns=" not in svg_str:
            svg_str = svg_str.replace("<svg", '<svg xmlns="http://www.w3.org/2000/svg"', 1)
        # BeautifulSoup html.parser 把属性全转小写，但 SVG 标准要求驼峰；
        # 不修复的话 resvg 取不到 viewBox/preserveAspectRatio，只渲染默认 0,0 区域 → 碎片
        svg_str = re.sub(r"\bviewbox=", "viewBox=", svg_str)
        svg_str = re.sub(r"\bpreserveaspectratio=", "preserveAspectRatio=", svg_str)
        svg_str = re.sub(r"\bxlink:href=", "xlink:href=", svg_str)  # html.parser 一般不动冒号
        # font_size=24 使 ex 单位对应合理像素大小；zoom=3 提升清晰度
        return svg_to_bytes(svg_str, font_size=24, zoom=3, dpi=0)
    except Exception as e:
        print(f"  [公式SVG渲染失败] {e}", file=sys.stderr)
        return None


def formula_to_png_bytes(latex_code: str) -> bytes | None:
    """将 LaTeX 代码渲染为透明背景 PNG，使用 matplotlib mathtext"""
    try:
        import os as _os
        import tempfile as _tmp
        # 避免 matplotlib 因 sandbox 下配置目录不可写而报错
        _os.environ.setdefault("MPLCONFIGDIR", _tmp.gettempdir())
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig = plt.figure(figsize=(0.01, 0.01))
        fig.patch.set_facecolor("none")
        text = r"$" + latex_code + r"$"
        t = fig.text(0, 0, text, fontsize=11)
        fig.canvas.draw()
        bbox = t.get_window_extent()
        dpi = 150
        w = max(bbox.width / dpi + 0.05, 0.1)
        h = max(bbox.height / dpi + 0.05, 0.1)
        fig.set_size_inches(w, h)
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                    facecolor="none", transparent=True, pad_inches=0.01)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"  [公式渲染失败] {latex_code[:50]}: {e}", file=sys.stderr)
        return None


def _color_str_to_rgb(color_str: str):
    """将 CSS 颜色字符串（#rrggbb / #rgb / rgb(r,g,b)）转为 (r,g,b) 元组，失败返回 None"""
    try:
        s = color_str.strip()
        if s.startswith("#"):
            h = s[1:]
            if len(h) == 3:
                h = h[0]*2 + h[1]*2 + h[2]*2
            if len(h) >= 6:
                return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        m = re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)', s)
        if m:
            return int(m.group(1)), int(m.group(2)), int(m.group(3))
    except Exception:
        pass
    return None


def _parse_css_colors(soup) -> dict:
    """从页面 <style> 标签提取 .class → 前景色字符串 映射（用于代码语法高亮等）"""
    color_map = {}
    for style_tag in soup.find_all("style"):
        css = style_tag.string or ""
        for m in re.finditer(r'\.(\w[\w-]*)[^{]*\{([^}]*)\}', css, re.S):
            cls = m.group(1)
            rules = m.group(2)
            # 只提取 color（不含 background-color）
            cm = re.search(r'(?<![a-z-])color\s*:\s*(#[0-9a-fA-F]{3,8}|rgb\([^)]+\))', rules)
            if cm:
                color_map[cls] = cm.group(1)
    return color_map


# CSS color inline style 提取正则
_CSS_COLOR_RE = re.compile(r'(?<![a-z-])color\s*:\s*(#[0-9a-fA-F]{3,8}|rgb\([^)]+\))', re.I)

# 微信 / mdnice 代码高亮类名 → 颜色（静态备用，CSS 外链无法抓取时使用）
# 参考 VS Code Light+ 主题，与微信编辑器默认高亮接近
_WECHAT_CODE_COLORS = {
    "code-snippet__keyword":    "#0000ff",
    "code-snippet__operator":   "#0000ff",
    "code-snippet__builtin":    "#0070c1",
    "code-snippet__built_in":   "#0070c1",
    "code-snippet__string":     "#a31515",
    "code-snippet__comment":    "#008000",
    "code-snippet__number":     "#098658",
    "code-snippet__literal":    "#098658",
    "code-snippet__class-name": "#267f99",
    "code-snippet__function":   "#795e26",
    "code-snippet__attribute":  "#001080",
    "code-snippet__property":   "#001080",
    "code-snippet__variable":   "#001080",
    "code-snippet__params":     "#001080",
    "code-snippet__tag":        "#800000",
    "code-snippet__type":       "#267f99",
    "code-snippet__symbol":     "#a31515",
    "code-snippet__selector":   "#800000",
    "code-snippet__meta":       "#808080",
    "code-snippet__deletion":   "#9a1212",
    "code-snippet__addition":   "#008000",
    # hljs 类名（部分文章使用 hljs 高亮器）
    "hljs-keyword":   "#0000ff",
    "hljs-string":    "#a31515",
    "hljs-comment":   "#008000",
    "hljs-number":    "#098658",
    "hljs-built_in":  "#0070c1",
    "hljs-attr":      "#001080",
    "hljs-function":  "#795e26",
    "hljs-title":     "#795e26",
    "hljs-type":      "#267f99",
    "hljs-meta":      "#808080",
    "hljs-literal":   "#098658",
}


def _add_hyperlink(p, text: str, url: str, color=None):
    """向段落添加一个带超链接的 run（Word OOXML 超链接）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    # 在 document relationships 里添加超链接关系
    part = p.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    # 构建 <w:hyperlink> XML 节点
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    # 添加 run
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # 超链接样式：蓝色下划线 + 正文字体
    _set_rfonts(rPr, _body_font())
    link_color = OxmlElement('w:color')
    if color:
        link_color.set(qn('w:val'), '%02X%02X%02X' % color)
    else:
        link_color.set(qn('w:val'), '0563C1')  # Word 默认超链接蓝
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    rPr.append(link_color)
    run.append(rPr)
    t = OxmlElement('w:t')
    import re as _re
    _t = _re.sub('(?<![0-9#*])\ufe0f|\ufe0f(?!\u20e3)', '', text)  # 仅保留 keycap 中的 FE0F
    _t = _re.sub('(?<![0-9#*\ufe0f])\u20e3', '', _t)            # 剥离非 keycap 的 U+20E3
    t.text = _t.replace('\ufe0e', '')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def _populate_para_with_colors(p, elem, css_colors=None, img_cache=None):
    """向已有段落添加带颜色的 run（支持行内公式、语法高亮色、内联 style 颜色、行内图片）"""
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor, Inches

    def _walk(node, color=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).replace("\xa0", " ")
                if text.strip():
                    # 去除行尾换行和空格（HTML 元素内多余空白）
                    cleaned = text.rstrip('\n\r ')
                    if not cleaned:
                        cleaned = text.strip()
                    # 剥离变体选择符，但保留合法 keycap 序列中的 FE0F
                    # 合法 keycap: [0-9#*] + U+FE0F + U+20E3（如 1️⃣ 5️⃣）
                    # 其他情况的 FE0F（如 ⚔️ 中的 FE0F）和非 keycap 的 U+20E3 均剥离，否则 Word 渲染为 □
                    import re as _re
                    cleaned = _re.sub('(?<![0-9#*])\ufe0f|\ufe0f(?!\u20e3)', '', cleaned)  # 仅保留 keycap 中的 FE0F
                    cleaned = _re.sub('(?<![0-9#*\ufe0f])\u20e3', '', cleaned)            # 剥离非 keycap 的 U+20E3
                    cleaned = cleaned.replace('\ufe0e', '')
                    # 按 emoji/非 emoji 分段
                    # emoji 段显式设 Segoe UI Emoji（Office 内置，Mac/Windows 均支持彩色 emoji）
                    # 若不设字体，继承 docDefaults 的 minorHAnsi（Calibri），Calibri 无 emoji 字形
                    for seg_is_emoji, seg in _split_text_by_emoji(cleaned):
                        if not seg:
                            continue
                        run = p.add_run(seg)
                        if seg_is_emoji and _is_keycap_seg(seg):
                            # keycap 序列（1️⃣ 4️⃣ 等）需要 Mac 原生 emoji 字体
                            # Segoe UI Emoji 在 Office for Mac 中不附带，设了等于无效
                            # Apple Color Emoji 是 macOS 系统 emoji 字体，Word for Mac 可用
                            run.font.name = KEYCAP_EMOJI_FONT
                        elif seg_is_emoji:
                            run.font.name = EMOJI_FONT
                        else:
                            _set_run_font(run, _body_font())
                        if color:
                            run.font.color.rgb = RGBColor(*color)
            elif isinstance(child, Tag):
                # 外层公式 span（cursor:pointer + data-formula）
                if (child.name == "span"
                        and child.has_attr("data-formula")
                        and _OUTER_FORMULA_RE.search(child.get("style", ""))):
                    latex = child["data-formula"]
                    png = formula_to_png_bytes(latex)
                    if png:
                        p.add_run().add_picture(BytesIO(png), height=Pt(_body_font_size()))
                    else:
                        p.add_run(f"[{latex}]")
                # 行内 MathJax 容器（无 display="true"，仅含 SVG）
                elif child.name == "mjx-container" and not child.get("display"):
                    png = mjx_svg_to_png(child)
                    if png:
                        p.add_run().add_picture(BytesIO(png), height=Pt(_body_font_size()))
                elif child.name == "br":
                    p.add_run("\n")
                elif child.name == "img":
                    # 行内/li 内图片：从 img_cache 取缓存字节插入
                    if img_cache is not None:
                        img_bytes = img_cache.get(id(child))
                        if img_bytes:
                            try:
                                p.add_run().add_picture(BytesIO(img_bytes), width=Inches(4))
                            except Exception:
                                pass
                elif child.name in ("table", "pre"):
                    pass  # 块级表格/代码块由主循环单独处理，避免在父级 section 中重复展开为纯文本
                elif child.name in ("svg", "mjx-container"):
                    pass  # SVG 和块级 mjx 由上层处理，不递归
                elif child.name in ("ol", "ul", "li"):
                    pass  # 嵌套列表由主循环处理，不在段落内展开
                elif child.name == "a" and child.get("href", "").startswith("http"):
                    # 超链接：提取颜色（如有），添加可点击链接
                    href = child["href"]
                    link_text = child.get_text(strip=False).replace("\xa0", " ")
                    if link_text.strip():
                        new_color = color
                        cm = _CSS_COLOR_RE.search(child.get("style", ""))
                        if cm:
                            extracted = _color_str_to_rgb(cm.group(1))
                            if extracted and not (extracted[0] > 240 and extracted[1] > 240 and extracted[2] > 240):
                                new_color = extracted
                        _add_hyperlink(p, link_text, href, new_color)
                else:
                    # 提取颜色：优先内联 style，其次 CSS 类
                    # 近白色（R/G/B 均 > 240）在 Word 白背景上不可见，直接丢弃
                    new_color = color
                    cm = _CSS_COLOR_RE.search(child.get("style", ""))
                    if cm:
                        extracted = _color_str_to_rgb(cm.group(1))
                        if extracted and not (extracted[0] > 240 and extracted[1] > 240 and extracted[2] > 240):
                            new_color = extracted
                    if new_color is color and css_colors:
                        for cls in (child.get("class") or []):
                            if cls in css_colors:
                                extracted = _color_str_to_rgb(css_colors[cls])
                                if extracted and not (extracted[0] > 240 and extracted[1] > 240 and extracted[2] > 240):
                                    new_color = extracted
                                break
                    _walk(child, new_color)

    _walk(elem)


def _add_rich_para(doc, elem, pt, align=None, css_colors=None, img_cache=None):
    """创建段落，将文本 run 和行内公式图片 run 混合排列，保留原始颜色"""
    p = doc.add_paragraph()
    _populate_para_with_colors(p, elem, css_colors, img_cache)
    _set_para_size(p, pt, align)
    return p


def _set_para_size(para, pt, align=None):
    """为段落设置字号、段前段后0、单倍行距，可选对齐方式"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # 段落级别默认字号（pPr/rPr）
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(pt * 2)))
    # 各 run 字号 + 字体
    # 注意：必须与 _split_text_by_emoji 逻辑一致，否则 keycap (1️⃣) 中的 '1'
    # 会被 _is_emoji_cp 误判为非 emoji，导致正文字体覆盖 emoji 字体
    for run in para.runs:
        run.font.size = Pt(pt)
        if not run.text:
            continue
        segs = _split_text_by_emoji(run.text)
        is_emoji_run = bool(segs) and all(is_e for is_e, _ in segs)
        if is_emoji_run and _is_keycap_seg(run.text):
            run.font.name = KEYCAP_EMOJI_FONT
        elif is_emoji_run:
            run.font.name = EMOJI_FONT
        else:
            _set_run_font(run, _body_font())
    # 段间距：段前段后0，单倍行距
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align is not None:
        para.alignment = align


def _add_code_para_colored(doc, pre_elem, css_colors=None):
    """将 <pre> 代码块转为带语法高亮颜色的 Word 段落"""
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor

    p = doc.add_paragraph()
    for s in doc.styles:
        if s.name == "No Spacing":
            p.style = s
            break

    def _add_run(text, color=None):
        if not text:
            return
        run = p.add_run(text)
        _set_run_font(run, _body_font())
        run.font.size = Pt(_body_font_size())
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _walk_code(node, color=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                _add_run(str(child).replace("\xa0", " "), color)
            elif isinstance(child, Tag):
                # <br> 在代码块内表示换行
                if child.name == "br":
                    _add_run("\n")
                    continue
                new_color = color
                # 内联 style
                cm = _CSS_COLOR_RE.search(child.get("style", ""))
                if cm:
                    new_color = _color_str_to_rgb(cm.group(1)) or color
                # CSS 类（代码高亮）：先查动态 CSS，再 fallback 到 _WECHAT_CODE_COLORS
                if new_color is color:
                    for cls in (child.get("class") or []):
                        if css_colors and cls in css_colors:
                            new_color = _color_str_to_rgb(css_colors[cls]) or color
                            break
                        if cls in _WECHAT_CODE_COLORS:
                            new_color = _color_str_to_rgb(_WECHAT_CODE_COLORS[cls]) or color
                            break
                _walk_code(child, new_color)

    code_children = pre_elem.find_all("code")
    if code_children:
        for i, code in enumerate(code_children):
            if i > 0:
                _add_run("\n")
            _walk_code(code)
    else:
        # 降级：无 <code> 子元素，直接按 <br> 分行
        for br in pre_elem.find_all("br"):
            br.replace_with("\n")
        _add_run(pre_elem.get_text().replace("\xa0", " ").strip("\n").rstrip())

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def build_docx(
        soup: BeautifulSoup,
        meta: dict,
        output: str,
        comments: list = None,
        body_font: str | None = None,
        body_font_size=None):
    font_token = _BODY_FONT_VAR.set(_normalize_body_font(body_font))
    size_token = _BODY_FONT_SIZE_VAR.set(_normalize_body_font_size(body_font_size))
    try:
        return _build_docx_impl(soup, meta, output, comments=comments)
    finally:
        _BODY_FONT_VAR.reset(font_token)
        _BODY_FONT_SIZE_VAR.reset(size_token)


def _build_docx_impl(soup: BeautifulSoup, meta: dict, output: str, comments: list = None):
    doc = Document()

    # 标题：去掉末尾标点，字号四号（14pt），居中，段前段后0
    clean_title = meta["title"].rstrip().rstrip("。！？，、；：…—,.!?;:")
    heading = doc.add_heading(clean_title, level=0)
    _set_para_size(heading, 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 元信息（默认五号 10.5pt，可由 Web UI 覆盖）
    info_parts = [p for p in [meta["account"], meta["author"], meta["publish_time"]] if p]
    if info_parts:
        p = doc.add_paragraph(" | ".join(info_parts))
        _set_para_size(p, _body_font_size(), align=WD_ALIGN_PARAGRAPH.CENTER)

    # 空行（段间距也清零）
    blank = doc.add_paragraph()
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(0)
    blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 正文：自动识别微信 / CSDN 容器
    content = soup.select_one("#js_content, .rich_media_content, #article_content, #content_views, .article_content")
    if not content:
        body_text = soup.body.get_text(strip=True) if soup.body else ""
        # 图片合集 / 沉浸式图文：#img_swiper 或 #img_list 存在，内容由 JS 渲染
        if soup.find(id="img_swiper") or soup.find(id="img_list"):
            fallback_text = _extract_meta_description_text(soup)
            if fallback_text:
                written = 0
                for block in re.split(r"\n{2,}", fallback_text):
                    block = "\n".join(line.strip() for line in block.splitlines()).strip()
                    if not block:
                        continue
                    p = doc.add_paragraph(block)
                    _set_para_size(p, _body_font_size())
                    written += 1
                note = doc.add_paragraph("注：此链接为微信图片合集/沉浸式图文，图片由微信客户端动态渲染；本文已提取页面内可用的文字正文。")
                _set_para_size(note, _body_font_size())
            else:
                p = doc.add_paragraph(
                    "此链接是微信图片合集/沉浸式图文。图片内容由微信客户端 JavaScript 动态渲染，"
                    "当前无法直接抓取图片；页面内也未提供可提取的文字正文。"
                )
                _set_para_size(p, _body_font_size())
                written = 0
            source_url = getattr(soup, '_url', '')
            if source_url:
                link_p = doc.add_paragraph("原文链接：")
                _add_hyperlink(link_p, source_url, source_url)
                _set_para_size(link_p, _body_font_size())
            doc.save(output)
            print(f"  正文：图片合集，提取文字 {written} 段 → {output}", file=sys.stderr)
            return
        elif soup.find("link", href=re.compile("secitptpage")) or soup.find(id="js_verify"):
            reason = "微信触发了访问验证（反爬虫）页面，请稍后重试，或检查链接是否带有多余的跟踪参数"
        elif "轻触阅读原文" in body_text or "阅读原文" in body_text:
            reason = "外链跳转卡片（内容在外部网站，无正文可抓取）"
        else:
            reason = "付费/受限/仅粉丝可见文章"
        print(f"[警告] 未找到正文容器：{reason}", file=sys.stderr)
        raise ValueError(f"无法提取正文：{reason}")

    # 提取页面 CSS 颜色映射（用于代码语法高亮、彩色文本等）
    css_colors = _parse_css_colors(soup)

    # 并发预下载所有图片
    img_cache: dict = {}
    img_elems = [(e, e.get("data-src") or e.get("src")) for e in content.find_all("img")]
    img_elems = [(e, u) for e, u in img_elems if u]
    # Referer：优先用 soup._url（fetch_article 存储的原始 URL），其次 og:url
    referer = getattr(soup, '_url', '')
    if not referer:
        og = soup.find("meta", property="og:url")
        if og:
            referer = og.get("content", "")
    if img_elems:
        print(f"  并发下载 {len(img_elems)} 张图片…", file=sys.stderr)
        with ThreadPoolExecutor(max_workers=10) as pool:
            future_map = {
                pool.submit(download_image, url, referer): elem
                for elem, url in img_elems
            }
            # 全局最多等 30s，超时直接放弃剩余图片
            done, not_done = futures_wait(future_map, timeout=30)
            for future in not_done:
                future.cancel()
                print("  [图片超时跳过]", file=sys.stderr)
            for future in done:
                elem = future_map[future]
                try:
                    img_cache[id(elem)] = future.result()
                except Exception:
                    pass

    img_count = 0
    para_count = 0
    skip_elems: set = set()  # 已由父元素处理、需跳过的元素 id

    for elem in content.descendants:
        if id(elem) in skip_elems:
            continue
        # 跳过 table 内部元素，避免重复（由 _add_table_to_doc 统一处理）
        if elem.find_parent("table"):
            continue
        # 跳过 pre 内部元素，避免重复（由 pre 分支统一处理）
        if elem.find_parent("pre"):
            continue
        # 跳过公式 span 内部元素（SVG 等，由 _add_rich_para 统一处理）
        if elem.find_parent("span", attrs={"data-formula": True}):
            continue
        # 跳过 mjx-container 内部元素（SVG，由各自分支统一处理）
        if elem.find_parent("mjx-container"):
            continue
        # 跳过 li 内部的非-li 元素（由 li 分支统一处理），嵌套 <li> 本身需处理
        if elem.find_parent("li") and elem.name != "li":
            continue
        # 块级 MathJax 公式（display=true），渲染为图片独占一行
        if elem.name == "mjx-container" and elem.get("display") == "true":
            png = mjx_svg_to_png(elem)
            if png:
                try:
                    doc.add_picture(BytesIO(png), width=Inches(5.5))
                    para_count += 1
                except Exception as e:
                    print(f"  [块级公式插入失败] {e}", file=sys.stderr)
            continue
        if elem.name == "table":
            if _is_code_table(elem):
                _add_code_table_to_doc(doc, elem, css_colors)
            else:
                _add_table_to_doc(doc, elem, _set_para_size)
            para_count += 1
        elif elem.name == "pre":
            # 代码块：使用 _add_code_para_colored 保留语法高亮颜色
            p = _add_code_para_colored(doc, elem, css_colors)
            # 检查是否有内容（纯空白 pre 跳过）
            if not p.text.strip():
                p._element.getparent().remove(p._element)
            else:
                para_count += 1
        elif elem.name == "img":
            img_bytes = img_cache.get(id(elem))
            if img_bytes:
                add_image_to_doc(doc, img_bytes)
                img_count += 1
        elif elem.name in ("p", "section", "h1", "h2", "h3", "h4", "h5", "h6", "li", "blockquote"):
            # li 列表项：不因有子元素而跳过，直接处理并生成序号/项目符号
            if elem.name == "li":
                if not elem.get_text(strip=True):
                    continue
                # 计算嵌套深度（每多一层 ol/ul 祖先 +1，顶层=0 不缩进）
                _depth = 0
                _node = elem.parent
                while _node and _node != content:
                    if _node.name in ("ol", "ul"):
                        _depth += 1
                    _node = _node.parent
                indent_level = max(0, _depth - 1)  # 顶层列表本身不缩进

                parent_list = elem.find_parent(["ol", "ul"])
                if parent_list and parent_list.name == "ol":
                    siblings = list(parent_list.find_all("li", recursive=False))
                    idx = (siblings.index(elem) + 1) if elem in siblings else 1
                    prefix = f"{idx}. "
                else:
                    # 不同层级用不同项目符号
                    _bullet_chars = ("•", "◦", "▪")
                    prefix = _bullet_chars[min(indent_level, len(_bullet_chars) - 1)] + " "
                p = doc.add_paragraph()
                prefix_run = p.add_run(prefix)
                _set_run_font(prefix_run, _body_font())
                _populate_para_with_colors(p, elem, css_colors, img_cache)
                _set_para_size(p, _body_font_size())
                # 应用缩进（每层 0.5 cm）
                if indent_level > 0:
                    from docx.shared import Cm
                    p.paragraph_format.left_indent = Cm(0.5 * indent_level)
                    p.paragraph_format.first_line_indent = Pt(0)
                from docx.oxml.ns import qn as _qn
                has_drawing = bool(p._p.findall('.//' + _qn('w:drawing')))
                if not p.text.strip() and not has_drawing:
                    p._element.getparent().remove(p._element)
                    continue
                para_count += 1
                continue
            # 只在叶子段落写文字，避免重复
            if elem.find(["p", "section", "li"]):
                continue

            # 检测行内公式：外层 span(data-formula+cursor:pointer) 或行内 mjx-container
            outer_formulas = elem.find_all(
                "span", attrs={"data-formula": True, "style": _OUTER_FORMULA_RE}
            )
            inline_mjx = elem.find_all("mjx-container", attrs={"display": False})
            # find_all 返回所有 mjx-container，再过滤掉 display="true" 的块级
            inline_mjx = [m for m in elem.find_all("mjx-container") if not m.get("display")]
            if outer_formulas or inline_mjx:
                # 将外层公式 span 及其子元素标记为已处理
                for fs in outer_formulas:
                    skip_elems.add(id(fs))
                    for d in fs.descendants:
                        skip_elems.add(id(d))
                # 将行内 mjx-container 及其子元素标记为已处理（_add_rich_para 会处理它们）
                for mj in inline_mjx:
                    skip_elems.add(id(mj))
                    for d in mj.descendants:
                        skip_elems.add(id(d))
                if elem.name in ("h1", "h2"):
                    p = _add_rich_para(doc, elem, 12, css_colors=css_colors, img_cache=img_cache)
                elif elem.name in ("h3", "h4"):
                    p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
                elif elem.name in ("h5", "h6"):
                    p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
                elif elem.name == "blockquote":
                    p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
                    try:
                        p.style = "Intense Quote"
                    except Exception:
                        pass
                else:
                    p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
                para_count += 1
                continue

            # 普通段落：先过滤无实质内容的空元素（空 <p>/<section> 等），
            # 与旧版 get_text(strip=True) 为空时 continue 的行为保持一致，
            # 避免创建后难以移除的空段落（30页→60页的根因）
            if not elem.get_text(strip=True):
                continue
            if elem.name in ("h1", "h2"):
                p = doc.add_heading("", level=2)
                _populate_para_with_colors(p, elem, css_colors, img_cache)
                _set_para_size(p, 12)  # 小四
            elif elem.name in ("h3", "h4"):
                p = doc.add_heading("", level=3)
                _populate_para_with_colors(p, elem, css_colors, img_cache)
                _set_para_size(p, _body_font_size())
            elif elem.name in ("h5", "h6"):
                p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
            elif elem.name == "blockquote":
                p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
                p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else p.style
            else:
                p = _add_rich_para(doc, elem, _body_font_size(), css_colors=css_colors, img_cache=img_cache)
            # 安全兜底：若段落仍为空（无文字且无行内图片），从文档移除
            from docx.oxml.ns import qn as _qn
            has_drawing = bool(p._p.findall('.//' + _qn('w:drawing')))
            if not p.text.strip() and not has_drawing:
                p._element.getparent().remove(p._element)
                continue
            para_count += 1

    # 统计实际图片数（含 li 内图片）
    from docx.oxml.ns import qn as _qn_final
    actual_img_count = len(doc.element.findall('.//' + _qn_final('w:drawing')))
    print(f"  正文：{para_count} 段落，{actual_img_count} 张图片", file=sys.stderr)

    # 评论
    if comments:
        doc.add_page_break()
        doc.add_heading("评论", level=1)
        for c in comments:
            nick = c.get("nick_name", "匿名")
            content_txt = c.get("content", "")
            like = c.get("like_num", 0)
            doc.add_paragraph(f"【{nick}】（赞 {like}）", style="Intense Quote" if False else None)
            doc.add_paragraph(content_txt)
            for reply in c.get("reply", {}).get("reply_list", []):
                doc.add_paragraph(
                    f"    ↳ {reply.get('nick_name','')}：{reply.get('content','')}"
                )
        print(f"  评论：{len(comments)} 条", file=sys.stderr)

    doc.save(output)
    print(f"  → {output}", file=sys.stderr)


# ============================================================
# 评论抓取（私有接口）
# ============================================================

def extract_comment_params(html: str) -> dict:
    """从文章 HTML 里抠出评论需要的参数"""
    params = {}
    for key in ("biz", "appmsgid", "comment_id", "idx", "sn"):
        m = re.search(rf'var\s+{key}\s*=\s*["\']([^"\']+)["\']', html)
        if m:
            params[key] = m.group(1)
    # __biz 在 URL / meta 里
    m = re.search(r'var\s+__biz\s*=\s*["\']([^"\']+)["\']', html)
    if m:
        params["__biz"] = m.group(1)
    return params


def fetch_comments(article_url: str) -> list:
    if not COMMENTS_COOKIE:
        print(
            "[评论] 未配置 COOKIES，跳过。请编辑脚本顶部 COMMENTS_COOKIE 变量。",
            file=sys.stderr,
        )
        return []

    headers = {
        "User-Agent": COMMENTS_USER_AGENT,
        "Cookie": COMMENTS_COOKIE,
        "Referer": article_url,
    }
    html = requests.get(article_url, headers=headers, timeout=20).text
    p = extract_comment_params(html)
    if not all(k in p for k in ("__biz", "appmsgid", "comment_id", "idx")):
        print(f"[评论] 无法解析参数：{p}", file=sys.stderr)
        return []

    api = "https://mp.weixin.qq.com/mp/appmsg_comment"
    qs = {
        "action": "getcomment",
        "scene": "0",
        "__biz": p["__biz"],
        "appmsgid": p["appmsgid"],
        "idx": p["idx"],
        "comment_id": p["comment_id"],
        "offset": "0",
        "limit": "100",
        "f": "json",
    }
    try:
        r = requests.get(api, params=qs, headers=headers, timeout=20)
        data = r.json()
        return data.get("elected_comment", []) + data.get("friend_comment", [])
    except Exception as e:
        print(f"[评论] 抓取失败：{e}", file=sys.stderr)
        return []


# ============================================================
# CLI
# ============================================================

def safe_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", name).strip()[:80] or "wechat_article"


def main():
    parser = argparse.ArgumentParser(description="微信公众号文章 → Word 文档")
    parser.add_argument("url", help="微信文章链接 https://mp.weixin.qq.com/s/...")
    parser.add_argument("-o", "--output", help="输出 docx 路径（默认用文章标题）")
    parser.add_argument("--font", choices=SUPPORTED_BODY_FONTS, default=DEFAULT_BODY_FONT,
                        help=f"正文字体（默认：{DEFAULT_BODY_FONT}）")
    parser.add_argument("--font-size", type=float, default=DEFAULT_BODY_FONT_SIZE_PT,
                        help=f"正文字号 pt（默认：{DEFAULT_BODY_FONT_SIZE_PT}）")
    parser.add_argument("--comments", action="store_true", help="抓取评论（需先配置 COOKIES）")
    args = parser.parse_args()

    print(f"抓取：{args.url}", file=sys.stderr)
    soup = fetch_article(args.url)
    meta = extract_meta(soup)
    print(f"标题：{meta['title']}", file=sys.stderr)

    output = args.output or f"{safe_filename(meta['title'])}.docx"

    comments = fetch_comments(args.url) if args.comments else None
    build_docx(soup, meta, output, comments=comments, body_font=args.font, body_font_size=args.font_size)


if __name__ == "__main__":
    main()
